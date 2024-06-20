# -*- coding: utf-8 -*-

from fastapi import APIRouter, HTTPException, UploadFile, status, BackgroundTasks
from fastapi.responses import StreamingResponse
from models.OCRModel import *
from models.RestfulModel import *
from paddleocr import PaddleOCR
from utils.ImageHelper import base64_to_ndarray, bytes_to_ndarray
import requests
import os
import io
from docx import Document
import shutil

OCR_LANGUAGE = os.environ.get("OCR_LANGUAGE", "en")

router = APIRouter(prefix="/ocr", tags=["OCR"])

ocr = PaddleOCR(use_angle_cls=True, lang=OCR_LANGUAGE)

def clear_temp_folder(folder_path):
    """
    Clears all files and folders within the specified folder path.
    """
    try:
        if os.path.exists(folder_path):  # Check if the folder exists
            shutil.rmtree(folder_path)
        os.makedirs(folder_path)  # Recreate the temp folder after clearing
    except Exception as e:
        print(f'Failed to clear {folder_path}. Reason: {e}')


# @router.get('/predict-by-path', response_model=RestfulModel, summary="Identify local images by path")
# def predict_by_path(image_path: str):
#     result = ocr.ocr(image_path, cls=True)
#     restfulModel = RestfulModel(
#         resultcode=200, message="Success", data=result, cls=OCRModel)
#     return restfulModel


# @router.post('/predict-by-base64', response_model=RestfulModel, summary="Recognize base64 data")
# def predict_by_base64(base64model: Base64PostModel):
#     img = base64_to_ndarray(base64model.base64_str)
#     result = ocr.ocr(img=img, cls=True)
#     restfulModel = RestfulModel(
#         resultcode=200, message="Success", data=result, cls=OCRModel)
#     return restfulModel


@router.post('/predict-by-file', summary="Identify uploaded files")
async def predict_by_file(file: UploadFile, background_tasks: BackgroundTasks):
    if file.filename.endswith((".jpg", ".png")):  # Only handle common format images
        file_data = await file.read()  # Read file asynchronously
        result = ocr.ocr(file_data, cls=True)  # Perform OCR on the image data

        # Convert OCR results to .docx
        extracted_text = ""
        for line in result:
            for word in line:
                extracted_text += word[1][0] + " "  # Extracted text is at index 0 in each word tuple

        doc = Document()
        doc.add_paragraph(extracted_text)
        
        # Save .docx file to BytesIO buffer
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)

        # Prepare filename for attachment
        filename_base = os.path.basename(file.filename).split('.')[0]

        # Return .docx file as a streaming response
        response = StreamingResponse(
            docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={'Content-Disposition': f'attachment; filename={filename_base}.docx'},
            status_code=status.HTTP_200_OK
        )
        background_tasks.add_task(clear_temp_folder, './temp')
        return response
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Please upload images in .jpg or .png format"
        )


@router.get('/predict-by-url', summary="Identify image URL")
async def predict_by_url(url: str, background_tasks: BackgroundTasks):
    if not url.lower().endswith((".jpg", ".png")):  # Only handle common format images
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Please provide URLs to images in .jpg or .png format"
        )

    urlresponse = requests.get(url)
    if urlresponse.status_code != 200:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unable to access the provided URL"
        )

    file_data = urlresponse.content
    result = ocr.ocr(file_data, cls=True)  # Perform OCR on the image data

    # Convert OCR results to .docx
    extracted_text = ""
    for line in result:
        for word in line:
            extracted_text += word[1][0] + " "  # Extracted text is at index 0 in each word tuple

    doc = Document()
    doc.add_paragraph(extracted_text)

    # Save .docx file to BytesIO buffer
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    # Prepare filename for attachment
    filename_base = os.path.basename(url).split('.')[0]

    # Return .docx file as a streaming response
    response = StreamingResponse(
        docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={'Content-Disposition': f'attachment; filename={filename_base}.docx'},
        status_code=status.HTTP_200_OK
    )
    background_tasks.add_task(clear_temp_folder, './temp')
    return response
